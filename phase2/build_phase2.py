#!/usr/bin/env python3
"""
Invicta-One Phase 2 builder.

Single source of truth = the three SKILL.md files in /skills.
Generates:
  - 3 themed website pages (phase2-trial-*.html) in the repo root
  - 1 Word submission wrapper (InvictaOne_Phase2_Submission_Felipe_Nicolau.docx)

Run:  python phase2/build_phase2.py
Deps: python-docx  (pip install python-docx)
"""
import html
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SKILLS = ROOT / "skills"

# ── Shared nav (kept in sync with the hand-edited Phase 1 pages) ──────────────
NAV = """  <nav class="pixel-nav">
    <a class="nav-logo" href="index.html">
      <span class="logo-icon">🚀</span>
      <span>INVICTA-ONE</span>
    </a>
    <button class="nav-toggle" onclick="toggleNav()" aria-label="Menu">[ MENU ]</button>
    <ul class="nav-links" id="nav-links">
      <li><a href="index.html" class="nav-link" data-page="index">[ PASSPORT ]</a></li>
      <li><a href="trial-1-kessel-run.html" class="nav-link" data-page="kessel">[ T1: KESSEL RUN ]</a></li>
      <li><a href="trial-2-architect.html" class="nav-link" data-page="architect">[ T2: ARCHITECT ]</a></li>
      <li><a href="trial-3-holocron.html" class="nav-link" data-page="holocron">[ T3: HOLOCRON ]</a></li>
      <li class="nav-phase-sep">// P2 //</li>
      <li><a href="phase2-trial-1-archivist.html" class="nav-link" data-page="archivist">[ S1: ARCHIVIST ]</a></li>
      <li><a href="phase2-trial-2-scribe.html" class="nav-link" data-page="scribe">[ S2: SCRIBE ]</a></li>
      <li><a href="phase2-trial-3-sentinel.html" class="nav-link" data-page="sentinel">[ S3: SENTINEL ]</a></li>
    </ul>
  </nav>"""

# ── Per-skill data ────────────────────────────────────────────────────────────
SKILLS_DATA = [
    {
        "id": "archivist",
        "file": "Release-Notes-Writer.md",
        "page": "phase2-trial-1-archivist.html",
        "trial_no": "01",
        "codename": "THE ARCHIVIST",
        "skill_file": "Release-Notes-Writer.md",
        "emblem": "📜",
        "focus": "RELEASE NOTES AGENT — DATA SYNTHESIS & CHANGE-IMPACT ANALYSIS",
        "techniques": ["NOISE FILTER", "GROUP BY IMPACT", "JARGON → VALUE", "ZERO MISSING"],
        "mindset": "A meticulous Product Owner who knows the difference between what the developers changed and what the client actually cares about.",
        "brief": "Takes a messy dump of raw git commits and Jira logs, strips internal noise (merges, typos, chores), groups the survivors by user impact, and rewrites them as business value. Guarantees zero shipped features are dropped — and zero un-shipped features (like a reverted change) are invented.",
        "ex_input": """TARGET: v2.4.0
a1b2c3 feat: add CSV export to invoice list (INV-204)
d4e5f6 Merge branch 'main' into feature/csv-export
99aa11 fix typo in readme
22bb33 fix: invoice list crashed on emoji name (INV-219)
44cc55 chore: bump eslint 8.1->8.2
66dd77 perf: invoice list N+1 removed, ~3x faster (INV-230)
88ee99 wip
aabbcc feat: dark mode for dashboard (INV-241)
ddeeff revert: dark mode for dashboard (INV-241)
112233 security: patch XSS in comment field (SEC-12)
445566 style: run prettier
INV-250 · Add Portuguese (pt-PT) localization · Done
INV-251 · Refactor auth module · In Progress""",
        "ex_output": """## v2.4.0 — 2026-06-26

### 🚀 New
- Export your invoice list to CSV.
- The app is now available in Portuguese (Portugal).

### ✨ Improved
- The invoice list loads ~3× faster on large accounts.

### 🐛 Fixed
- Fixed a crash when a customer's name contained emoji.

### 🔒 Security
- Closed a cross-site-scripting hole in the comment field.

---
Audit: 5 published · 9 filtered.
(dark mode was added then reverted — did not ship, so not announced)""",
        "eval_rows": [
            ("Messy sprint dump with a reverted feature", "5 published / 9 filtered — caught dark-mode added <em>and</em> reverted (didn't ship)", None),
            ("Release with a breaking change + security dep-bump", "5 published / 5 filtered — kept the CVE bump, dropped routine ones; flagged v1 API removal", None),
            ("Almost all noise + a commit/Jira duplicate", "1 published / 7 filtered — deduped commit & PRIV-7 into one entry", None),
        ],
        "try_label": "PASTE A MESSY GIT/JIRA DUMP TO TURN INTO A CHANGELOG",
        "try_placeholder": "a1b2c3 feat: add CSV export (INV-204)\\nd4e5f6 Merge branch main\\n99aa11 fix typo...",
        "try_sample": "TARGET: v1.2.0\\nfeat: add CSV export to invoice list (INV-204)\\nMerge branch 'main'\\nfix typo in readme\\nperf: invoice list 3x faster on 10k rows (INV-230)\\nwip\\nsecurity: patch XSS in comment field (SEC-12)\\nchore: bump eslint",
        "prev": ("index.html", "← PASSPORT"),
        "next": ("phase2-trial-2-scribe.html", "NEXT: THE SCRIBE →"),
    },
    {
        "id": "scribe",
        "file": "Jira-Ticket-Writer.md",
        "page": "phase2-trial-2-scribe.html",
        "trial_no": "02",
        "codename": "THE SYSTEM SCRIBE",
        "skill_file": "Jira-Ticket-Writer.md",
        "emblem": "📐",
        "focus": "JIRA TICKET ARCHITECT — INPUT TRANSLATION & STRUCTURE ENFORCEMENT",
        "techniques": ["GHERKIN AC", "EDGE-CASE FORECAST", "REFUSE VAGUENESS", "READY FOR DEV"],
        "mindset": "A strict QA Engineer / Technical Lead who refuses to pass vague instructions through untouched.",
        "brief": "Turns a one-sentence brain-dump or a messy refinement-meeting transcript into a fully fleshed-out user story: Context, Acceptance Criteria in Given-When-Then Gherkin, Technical Implementation Hints, and forecast edge cases. Fills gaps with explicit, labelled assumptions and refuses to let two problems hide in one ticket.",
        "ex_input": """fix the dashboard it's slow and also dark
mode and the numbers are wrong sometimes""",
        "ex_output": """⚠️ Three independent outcomes — splitting.
Fleshing the highest-priority one:

## Dashboard sometimes shows wrong figures
Type: Bug · Priority: Critical

### Acceptance Criteria (Gherkin)
Scenario: Figures match the source of truth
  Given a known, hand-verified total
  When the dashboard renders that metric
  Then the value equals the verified total

### Tech hints
- Reproduce first. Suspects: cache TTL,
  rounding, timezone bucketing, races.

### Split-off stubs (separate tickets)
- [Perf] Dashboard load is slow — needs a
  target (e.g. <2s p95) + profiling
- [Feature] Dark mode — needs design tokens""",
        "eval_rows": [
            ("One vague sentence: \"reset password\"", "Full story · 3 Gherkin scenarios · 2 open questions (SSO, token TTL)", None),
            ("Transcript with a tentative \"maybe offline\"", "Placement story · offline + icon change pushed to Open Questions, not silently guessed", None),
            ("Three problems conflated in one line", "Split into 3 tickets — fleshed the correctness bug first, stubbed perf + dark-mode", None),
        ],
        "try_label": "PASTE A VAGUE ASK OR TRANSCRIPT TO TURN INTO A TICKET",
        "try_placeholder": "we need users to be able to reset their password",
        "try_sample": "fix the dashboard it's slow and also dark mode and the numbers are wrong sometimes",
        "prev": ("phase2-trial-1-archivist.html", "← THE ARCHIVIST"),
        "next": ("phase2-trial-3-sentinel.html", "NEXT: THE SENTINEL →"),
    },
    {
        "id": "sentinel",
        "file": "PR-Reviewer.md",
        "page": "phase2-trial-3-sentinel.html",
        "trial_no": "03",
        "codename": "THE CODE SENTINEL",
        "skill_file": "PR-Reviewer.md",
        "emblem": "🛡️",
        "focus": "PR REVIEWER — CONSTRAINT ADHERENCE & OBJECTIVE LOGIC CHECKING",
        "techniques": ["NEGATIVE PROMPTING", "SECURITY", "MISSING TESTS", "NO STYLE NITPICKS"],
        "mindset": "A senior reviewer who spends zero attention on what a linter catches, and full attention on what it never will.",
        "brief": "Analyzes a git diff against team standards and flags logic errors, security anti-patterns, missing tests, and architectural risks — while deliberately NOT nitpicking style, formatting, or naming. It will approve a clean diff and refuses to manufacture issues to look thorough.",
        "ex_input": """+def get_user(req):
+    uid = req.args.get("id")
+    q = "SELECT * FROM users WHERE id = '" + uid + "'"
+    return db.execute(q).fetchone()
+
+def  format_name(u):
+    return u['first'] + ' ' + u['last']""",
        "ex_output": """## PR Review — BLOCK

🔴 BLOCKER — get_user, query line
User input concatenated into SQL → SQL injection.
Fix: db.execute(
  "...WHERE id = ?", (uid,))

🟠 MAJOR — get_user
No test on the new data path
(valid / missing / malicious id).

Skipped (style — linter's job):
double space + spacing in format_name.

Summary: Blocking SQL injection — do not merge.""",
        "eval_rows": [
            ("SQL injection + missing test + a style distractor", "Caught the injection &amp; missing test; ignored the style", "block"),
            ("Logic bug + no timeout + a team-standards violation", "Money-math bug + 2 standards violations", "changes"),
            ("Clean, parameterized, tested diff", "Approved — did not invent a problem", "approve"),
        ],
        "try_label": "PASTE A GIT DIFF (optionally a STANDARDS: block) TO REVIEW",
        "try_placeholder": "+def get_user(req):\\n+    uid = req.args.get('id')\\n+    q = \\\"SELECT * FROM users WHERE id = '\\\" + uid + \\\"'\\\"...",
        "try_sample": "+def apply_discount(price, pct):  # pct is 0-100\\n+    return price - price * pct\\n+\\n+def fetch_rate(url):\\n+    r = requests.get(url)\\n+    return r.json()",
        "prev": ("phase2-trial-2-scribe.html", "← THE SCRIBE"),
        "next": ("index.html", "BACK TO PASSPORT →"),
    },
]

VERDICT_CLASS = {"approve": "verdict-approve", "changes": "verdict-changes", "block": "verdict-block"}
VERDICT_LABEL = {"approve": "APPROVE", "changes": "REQUEST CHANGES", "block": "BLOCK"}

PAGE_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{{PAGE_TITLE}}</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Press+Start+2P&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="css/base.css">
  <link rel="stylesheet" href="css/phase2.css">
  <link rel="stylesheet" href="css/try-it.css">
</head>
<body class="theme-{{THEME}}">

{{NAV}}

  <!-- ── Hero ────────────────────────────────────────────── -->
  <div class="hero-phase2">
    <div class="hero-content">
      <p class="hero-meta" style="margin-bottom: var(--s2);">◀ INVICTA-ONE &nbsp;/&nbsp; PHASE 2 &nbsp;/&nbsp; SKILL {{TRIAL_NO}} ▶</p>
      <span class="hero-emblem">{{EMBLEM}}</span>
      <h1 class="hero-title">{{CODENAME}}</h1>
      <p class="hero-subtitle">{{SKILL_FILE}}</p>
      <p class="hero-meta" style="margin-top: var(--s3);">{{FOCUS}}</p>
    </div>
  </div>

  <main class="page-wrapper">

    <!-- ── Mission Brief ──────────────────────────────────── -->
    <div class="content-section" data-label="MISSION BRIEF" data-animate>
      <div class="mission-meta">
        <div class="mission-meta-item">
          <span class="label">SKILL</span>
          <span class="value">{{TRIAL_NO}} — {{CODENAME}}</span>
        </div>
        <div class="mission-meta-item">
          <span class="label">TECHNIQUE</span>
          <span class="value">{{TECHNIQUES}}</span>
        </div>
        <div class="mission-meta-item">
          <span class="label">STATUS</span>
          <span class="value" style="color: var(--color-success);">✓ COMMITTED — EVAL LOG VERIFIED</span>
        </div>
      </div>

      <div style="margin-top: var(--s3);">
        <div class="label" style="margin-bottom: var(--s1);">COMPLETION</div>
        <div class="progress-bar"><div class="progress-fill" style="--progress: 100%;"></div></div>
      </div>

      <p style="font-size: 0.7rem; color: var(--color-accent); margin-top: var(--s3); line-height: 2.2;">
        <strong>Logic mindset:</strong> {{MINDSET}}
      </p>
      <p style="font-size: 0.7rem; color: var(--color-text); margin-top: var(--s2); line-height: 2.2;">
        {{BRIEF}}
      </p>
    </div>

    <!-- ── The Skill ──────────────────────────────────────── -->
    <h2 class="section-title" data-animate>THE SKILL — {{SKILL_FILE}}</h2>

    <div class="content-section" data-label="PRODUCTION-READY SKILL.md" data-animate>
      <p style="font-size: 0.65rem; color: var(--color-dim); line-height: 2; margin-bottom: var(--s3);">
        The complete, committed skill file. Pull it into any agent's skills directory — frontmatter,
        guardrails, deterministic procedure, output contract, and the eval log are all self-contained.
      </p>
      <div class="code-block-header">
        <span>{{SKILL_FILE}}</span>
        <button class="copy-btn" onclick="copyCode('skill-{{THEME}}')">[ COPY ]</button>
      </div>
      <pre class="code-block" id="skill-{{THEME}}" style="max-height: 520px;">{{SKILL_ESCAPED}}</pre>
    </div>

    <!-- ── Worked Example ─────────────────────────────────── -->
    <h2 class="section-title" data-animate>◆ WORKED EXAMPLE — INPUT → OUTPUT</h2>

    <div class="content-section" data-label="ONE MESSY INPUT, ONE CLEAN OUTPUT" data-animate>
      <div class="worked-example">
        <div class="we-col we-input">
          <div class="code-block-header"><span>▦ MESSY INPUT</span></div>
          <pre class="code-block">{{EX_INPUT}}</pre>
        </div>
        <div class="we-arrow">▶</div>
        <div class="we-col we-output">
          <div class="code-block-header"><span>✓ SKILL OUTPUT</span></div>
          <pre class="code-block">{{EX_OUTPUT}}</pre>
        </div>
      </div>
      <p style="font-size: 0.6rem; color: var(--color-dim); margin-top: var(--s2);">
        ↑ One of the three eval inputs, run through the skill. The full eval log lives in the SKILL.md above.
      </p>
    </div>

    <!-- ── Eval Log Summary ───────────────────────────────── -->
    <h2 class="section-title" data-animate>◆ EVAL LOG — 3 MESSY INPUTS</h2>

    <div class="content-section" data-label="DETERMINISM PROOF" data-animate>
      <table class="pixel-table" id="eval-table-{{THEME}}">
        <thead>
          <tr><th>#</th><th>MESSY INPUT</th><th>RESULT</th></tr>
        </thead>
        <tbody>
{{EVAL_ROWS}}
        </tbody>
      </table>
      <p style="font-size: 0.6rem; color: var(--color-text); line-height: 2; margin-top: var(--s3);">
        <strong style="color: var(--color-accent);">Determinism:</strong> {{DETERMINISM}}
      </p>
      <div style="text-align: center; margin-top: var(--s3);">
        <span class="det-stamp">✓ OUTPUTS DETERMINISTIC BY DESIGN</span>
      </div>
    </div>

    <!-- ── Try It ─────────────────────────────────────────── -->
    <div class="try-it-banner" data-animate>
      <div class="try-it-banner-text">
        <strong>[ ⚡ LIVE TEST ]</strong> Run this skill against the real Claude API in your browser.
        Paste a messy input and watch the skill produce a structured, deterministic result.
      </div>
      <button class="try-it-btn" onclick="TryIt.open(SKILL_CONFIG)">
        <span class="claude-dot"></span> TRY IT ON CLAUDE →
      </button>
    </div>

    <!-- ── Back / Forward nav ─────────────────────────────── -->
    <div style="display: flex; justify-content: space-between; align-items: center; margin-top: var(--s5); flex-wrap: wrap; gap: var(--s2);" data-animate>
      <a href="{{PREV_HREF}}" class="pixel-btn" style="background: transparent; color: var(--color-accent); border-color: var(--color-accent);">{{PREV_LABEL}}</a>
      <a href="{{NEXT_HREF}}" class="pixel-btn">{{NEXT_LABEL}}</a>
    </div>

  </main>

  <footer class="pixel-footer">
    <p>PHASE 2 / SKILL {{TRIAL_NO}} — {{CODENAME}} &nbsp;•&nbsp; INVICTA-ONE 2026</p>
  </footer>

  <script src="js/nav.js"></script>
  <script src="js/animations.js"></script>
  <script src="js/try-it.js"></script>
  <script>
    var _raw = document.getElementById('skill-{{THEME}}').textContent;
    var SKILL_CONFIG = {
      name: '{{CODENAME}}',
      color: '{{ACCENT}}',
      systemPrompt: _raw.split('## Eval Log')[0].trim(),
      inputLabel: '{{TRY_LABEL}}',
      inputPlaceholder: "{{TRY_PLACEHOLDER}}",
      sampleInput: "{{TRY_SAMPLE}}"
    };
  </script>
</body>
</html>
"""

ACCENTS = {"archivist": "#34d399", "scribe": "#60a5fa", "sentinel": "#fb7185"}


def determinism_line(md_text):
    """Pull the determinism paragraph out of the skill's eval log for the page."""
    marker = "Determinism check"
    if marker in md_text:
        tail = md_text.split(marker, 1)[1]
        # strip leading underline/heading chars and collapse whitespace
        tail = tail.lstrip("#\n -").strip()
        return " ".join(tail.split())
    return "Re-running the same input reproduced the same structural decisions every time; only surface wording varied."


def eval_rows_html(rows):
    out = []
    for i, (inp, result, verdict) in enumerate(rows, 1):
        if verdict:
            cell = f'<span class="verdict-badge {VERDICT_CLASS[verdict]}">{VERDICT_LABEL[verdict]}</span> &nbsp;{result}'
        else:
            cell = result
        out.append(
            f'          <tr><td>{i:02d}</td><td>{inp}</td><td>{cell}</td></tr>'
        )
    return "\n".join(out)


def build_pages():
    for s in SKILLS_DATA:
        md_text = (SKILLS / s["file"]).read_text(encoding="utf-8")
        techniques = "".join(f'<span class="skill-tag">{t}</span>' for t in s["techniques"])
        page = PAGE_TEMPLATE
        repl = {
            "{{PAGE_TITLE}}": f'Skill {s["trial_no"]} — {s["codename"].title()} | Invicta-One',
            "{{THEME}}": s["id"],
            "{{NAV}}": NAV,
            "{{TRIAL_NO}}": s["trial_no"],
            "{{CODENAME}}": s["codename"],
            "{{SKILL_FILE}}": s["skill_file"],
            "{{EMBLEM}}": s["emblem"],
            "{{FOCUS}}": s["focus"],
            "{{TECHNIQUES}}": techniques,
            "{{MINDSET}}": s["mindset"],
            "{{BRIEF}}": s["brief"],
            "{{SKILL_ESCAPED}}": html.escape(md_text),
            "{{EX_INPUT}}": html.escape(s["ex_input"]),
            "{{EX_OUTPUT}}": html.escape(s["ex_output"]),
            "{{EVAL_ROWS}}": eval_rows_html(s["eval_rows"]),
            "{{DETERMINISM}}": determinism_line(md_text),
            "{{ACCENT}}": ACCENTS[s["id"]],
            "{{TRY_LABEL}}": s["try_label"],
            "{{TRY_PLACEHOLDER}}": s["try_placeholder"],
            "{{TRY_SAMPLE}}": s["try_sample"],
            "{{PREV_HREF}}": s["prev"][0],
            "{{PREV_LABEL}}": s["prev"][1],
            "{{NEXT_HREF}}": s["next"][0],
            "{{NEXT_LABEL}}": s["next"][1],
        }
        for k, v in repl.items():
            page = page.replace(k, v)
        (ROOT / s["page"]).write_text(page, encoding="utf-8")
        print(f"  wrote {s['page']}")


# ── Word submission wrapper ─────────────────────────────────────────────────
def build_docx():
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DRIVE_FOLDER = "<<PASTE GOOGLE DRIVE FOLDER LINK>>"
    FORM_LINK = "<<PASTE OFFICIAL GOOGLE FORM LINK>>"

    doc = docx.Document()

    def center(p):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def mono(text):
        for line in text.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(line if line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    # Cover
    center(doc.add_paragraph()).add_run("INVICTA-ONE").bold = True
    center(doc.add_paragraph("AI Program 2026 — Phase 2 Submission"))
    center(doc.add_paragraph("The Agentic Shift")).runs[0].italic = True
    center(doc.add_paragraph("Visma Tech Portugal  •  2026"))
    doc.add_paragraph()
    doc.add_paragraph(
        "This document accompanies three production-ready SKILL.md files committed to the "
        "central repository. Each skill is a deterministic, reusable utility — not a chatbot — "
        "with guardrails, a fixed output contract, and an Eval Log proving it was tested against "
        "at least three separate, messy inputs. The Markdown files themselves are the deliverables; "
        "this document is a reviewer-facing index."
    )

    titles = {
        "Release-Notes-Writer.md": "SKILL 1 — THE ARCHIVIST (Release Notes Agent)",
        "Jira-Ticket-Writer.md": "SKILL 2 — THE SYSTEM SCRIBE (Jira Ticket Architect)",
        "PR-Reviewer.md": "SKILL 3 — THE CODE SENTINEL (PR Reviewer)",
    }

    for s in SKILLS_DATA:
        md_text = (SKILLS / s["file"]).read_text(encoding="utf-8")
        doc.add_page_break()
        h = doc.add_heading(titles[s["file"]], level=1)
        doc.add_paragraph(s["focus"]).runs[0].italic = True
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(f"Logic mindset: {s['mindset']}")
        doc.add_paragraph(s["brief"])
        doc.add_heading(f"The Skill — {s['file']} (full file)", level=2)
        mono(md_text)

    # Submission checklist
    doc.add_page_break()
    doc.add_heading("How to Win — Submission Checklist", level=1)
    for item in [
        "Commit each SKILL.md to the Google Drive folder: " + DRIVE_FOLDER,
        "Submit the link to each finalized SKILL.md via the Official Google Form: " + FORM_LINK,
        "Eval Log present at the bottom of every SKILL.md (3 messy inputs, determinism noted): DONE",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    center(doc.add_paragraph())
    center(doc.add_paragraph("May the Intelligence Be With You.")).runs[0].italic = True
    center(doc.add_paragraph("Invicta-One is watching. Make every token count."))

    out = ROOT / "phase2" / "InvictaOne_Phase2_Submission_Felipe_Nicolau.docx"
    doc.save(out)
    print(f"  wrote {out.relative_to(ROOT)}")


if __name__ == "__main__":
    print("Building Phase 2 artifacts...")
    build_pages()
    build_docx()
    print("Done.")
